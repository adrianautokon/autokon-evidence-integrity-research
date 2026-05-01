"""
AutoKon Evidence Integrity R&D — DOCX builder
Reads: EVIDENCE_INTEGRITY_RESEARCH.md (the master MD)
Writes: AutoKon-Evidence-Integrity-Research.docx (full traditional report)
Run: python3 build-docx.py
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# AutoKon brand colors
SLATE_900 = RGBColor(0x0F, 0x17, 0x2A)
SLATE_700 = RGBColor(0x33, 0x41, 0x55)
SLATE_600 = RGBColor(0x47, 0x55, 0x69)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
SLATE_400 = RGBColor(0x94, 0xA3, 0xB8)
SLATE_200 = RGBColor(0xE2, 0xE8, 0xF0)
SLATE_100 = RGBColor(0xF1, 0xF5, 0xF9)
BLUE_600 = RGBColor(0x25, 0x63, 0xEB)
BLUE_700 = RGBColor(0x1D, 0x4E, 0xD8)
ORANGE = RGBColor(0xFF, 0x74, 0x26)
ORANGE_DARK = RGBColor(0xC2, 0x41, 0x0C)
GREEN = RGBColor(0x16, 0x65, 0x34)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_styled_para(doc, text, *, bold=False, size=11, color=SLATE_700, italic=False, align=None, family="Inter"):
    p = doc.add_paragraph()
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = family
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 22, 2: 17, 3: 13, 4: 11}
    colors = {1: SLATE_900, 2: SLATE_900, 3: SLATE_900, 4: SLATE_700}
    spaces_before = {1: 24, 2: 20, 3: 14, 4: 10}
    spaces_after = {1: 12, 2: 8, 3: 6, 4: 4}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before[level])
    p.paragraph_format.space_after = Pt(spaces_after[level])
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Poppins"
    run.font.size = Pt(sizes[level])
    run.font.color.rgb = colors[level]
    return p

def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "JetBrains Mono"
    run.font.size = Pt(8.5)
    run.font.color.rgb = SLATE_500

def add_callout(doc, text, *, color="slate"):
    color_map = {
        "slate": SLATE_100,
        "blue": RGBColor(0xEF, 0xF6, 0xFF),
        "orange": RGBColor(0xFF, 0xF7, 0xED),
        "green": RGBColor(0xDC, 0xFC, 0xE7),
    }
    color_hex_map = {
        "slate": "F1F5F9",
        "blue": "EFF6FF",
        "orange": "FFF7ED",
        "green": "DCFCE7",
    }
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(10.5)
    run.font.color.rgb = SLATE_700
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid"
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, "F1F5F9")
        for para in cell.paragraphs:
            run = para.runs[0] if para.runs else para.add_run(h)
            if para.runs:
                run.text = h
            run.bold = True
            run.font.name = "JetBrains Mono"
            run.font.size = Pt(8.5)
            run.font.color.rgb = SLATE_500
            if not para.runs:
                pass
    # Rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell_val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            for para in cell.paragraphs:
                run = para.add_run(str(cell_val))
                run.font.name = "Inter"
                run.font.size = Pt(9.5)
                run.font.color.rgb = SLATE_700
    # Apply column widths if given
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_bullet(doc, text, *, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(10.5)
    run.font.color.rgb = SLATE_700

def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(10.5)
    run.font.color.rgb = SLATE_700

# ===========================================
# Build the document
# ===========================================

doc = Document()

# Document margins
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

# === Cover ===
add_styled_para(doc, "EVIDENCE INTEGRITY R&D BRIEF · V1.0 · 2026-05-01",
                size=8.5, color=ORANGE_DARK, bold=True, family="JetBrains Mono")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("AutoKon — Evidence Integrity R&D Brief")
run.bold = True
run.font.name = "Poppins"
run.font.size = Pt(28)
run.font.color.rgb = SLATE_900

add_styled_para(doc, "Six Pillars to make AutoKon's evidence stand up to bank scrutiny.",
                size=14, color=SLATE_600, italic=True)

doc.add_paragraph()  # spacer

meta = [
    ("Authors", "Adrian (VP Product, AutoKon) + Claude Cowork"),
    ("Builds on", "Dickson + Claude evidence integrity summary (May 1, 2026)"),
    ("For review by", "Dickson (CEO), Faishal (CTO)"),
    ("Status", "Comprehensive R&D brief — disregards engineering / operational difficulty"),
    ("Version", "v1.0 · 2026-05-01"),
]
for k, v in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(k.upper() + "  ")
    r1.bold = True
    r1.font.name = "JetBrains Mono"
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = SLATE_500
    r2 = p.add_run(v)
    r2.font.name = "Inter"
    r2.font.size = Pt(11)
    r2.font.color.rgb = SLATE_700

doc.add_page_break()

# === TL;DR ===
add_heading(doc, "TL;DR", 1)
tldr_paras = [
    "The bank pitch (BSN/BTN) hinges on one promise: \"prove the photo is real, taken now, of the unit it claims to be — and the audit trail behind it cannot be tampered with.\" Dickson's prior work landed a 10-layer additive validation pipeline organized by engineering layer; this brief reorganizes that work — plus selectively-revived rejected items, plus new categories that close the bank-pitch gap — into 6 audience-first pillars that map to bank concerns directly.",
    "The 6 pillars are: Authenticity (is this photo real?), Identity (is the submitter who they claim?), Place (is this the actual unit?), Time (was this captured now?), Chain of Custody (can this record be tampered with post-hoc?), and Independent Verification (do you ever rely on something other than your own software?).",
    "Across 6 pillars there are 35 items, with 100+ proposed solutions analyzed across multiple dimensions: cost, engineering complexity, implementability today, low-end-device friendliness, Pelaksana friction, pitch story strength, false-positive risk, and adversarial difficulty. Two cross-cutting lenses introduced during the R&D brainstorm — threat motivation and attacker tech-savviness — shape the recommended ship priorities.",
    "Two key takeaways for the bank pitch: (1) AutoKon is engineered for Indonesian field reality — equatorial sun-angle math, BMKG weather feeds, call-to-prayer audio cross-check, RT/RW witness signatures, low-tech consumer-app fraud detectors, siteplan-relative GPS pattern matching that respects archipelago infrastructure realities. No US/EU competitor would think to build any of these. (2) AutoKon makes the inspector 10x more focused — software pre-screens, KJPP physical inspection inspects what software flags. ~80% of realistic fraud volume caught deterministically; residual 20% is structurally chain-of-custody.",
]
for tp in tldr_paras:
    add_styled_para(doc, tp, size=11, color=SLATE_700)

doc.add_page_break()

# === Section 1: Locked Principles ===
add_heading(doc, "§1 Locked Principles", 1)
add_styled_para(doc, "These are inherited from the AutoKon reports-roadmap PRD v3.3 (OpusPrime + Dickson, April 2026) and the Pak Jun product recommendations. This brief respects them without modification.")

add_heading(doc, "1.1 Quality is out of scope", 2)
add_styled_para(doc, "AutoKon verifies that progress was done, not that the work is high quality. Out: SMKK compliance, QC inspections, warranty defect audits. In: progress verification, rework tracking (financial ledger), retention release gates.")

add_heading(doc, "1.2 Developers are the customer. Contractors are the counterparty.", 2)
add_styled_para(doc, "Every AutoKon user works for the developer. Contractors never touch AutoKon. PDF signatures: developer-org roles only. Document voice: \"developer certifies\" not \"both parties agree.\"")

add_heading(doc, "1.3 The bank is the primary user (Pak Jun's reframe)", 2)
add_styled_para(doc, "For the bank pitch context, the bank is the buyer. Reports go to the bank as the audience; raw photos and process noise stay internal. Reports, not raw process, are the bank-facing surface.")

add_heading(doc, "1.4 No recurring ground discipline", 2)
add_styled_para(doc, "One-time-cost items (KTP enrollment, siteplan upload, 360° baseline, drone overflight if developer already operates one) are eligible for revival. Recurring discipline stays rejected.")

add_heading(doc, "1.5 Indonesian field reality is a binding constraint", 2)
add_bullet(doc, "Cell-tower triangulation in eastern Indonesia is poor. GPS readings drift ±30m+ in low-coverage provinces (NTT, Maluku, Papua).")
add_bullet(doc, "Indonesians use full addresses, not pin coordinates. Shopee, Tokopedia, Grab/Gojek default to typed addresses + landmarks.")
add_bullet(doc, "Most Indonesian field fraudsters are low-tech with high motivation. Microsoft Word export, Instagram Story overlays, screenshots, photo-of-printed-photo. Item 1.6 explicitly addresses these.")

doc.add_page_break()

# === Section 2: Two Cross-Cutting Lenses ===
add_heading(doc, "§2 Two Cross-Cutting Lenses", 1)

add_heading(doc, "2.1 Threat Motivation", 2)
add_styled_para(doc, "For each fraud type, who has the economic motivation to commit it, and how strongly?")
add_table(doc,
    headers=["Threat actor", "Strong motivation", "Weak motivation"],
    rows=[
        ("Developer (primary)", "Recycling photos · Adjacent-unit substitution · Backdating by weeks · AI-generated photos · Photo-of-printed-photo", "Sub-hour time-shift · Day-of-week · Cryptographic forgery"),
        ("Pelaksana (secondary)", "Buddy submission · Sloppy backdating · Account share", "Edit-forgery · AI-generation"),
        ("External attacker", "All of the above", "Volume-driven (low payoff)"),
        ("Insider", "Post-hoc DB tampering · Audit-log forgery · Signature backdating", "Real-time submission spoofing"),
    ],
    col_widths=[1.5, 3.0, 2.0])

doc.add_paragraph()

add_heading(doc, "2.2 Attacker Tech-Savviness", 2)
add_table(doc,
    headers=["Profile", "Common attacks", "Coverage"],
    rows=[
        ("Low-tech (most common)", "Word PDF export · IG Story overlays · Screenshots · Photo-of-print · WhatsApp re-share", "Pillars 1-4 catch ~80%. Item 1.6 adds explicit detectors."),
        ("Medium-tech", "EXIF manipulation · Custom timestamp-camera spoofs · Fake-GPS apps", "Pillars 1, 2, 3. Identity binding primary."),
        ("High-tech (very rare)", "AI-generated photos · PRNU manipulation · Cryptographic forgery · Insider attacks", "Pillar 1.4 + Pillar 5. Banks audit even though volume low."),
    ],
    col_widths=[1.5, 3.0, 2.0])

doc.add_paragraph()

add_heading(doc, "2.3 The 2x2", 2)
add_styled_para(doc, "The volume-of-fraud quadrant is bottom-right (low-tech × high motivation) — that's where AutoKon must focus engineering effort.", italic=True, color=SLATE_600)

doc.add_page_break()

# === Section 3: The 6 Pillars ===
add_heading(doc, "§3 The Six Pillars", 1)
add_styled_para(doc, "Each pillar opens with the bank's question, names its scope, and walks each item with: threat description, threat motivation, Dickson's existing solution if any, 2-5 proposed solutions with multi-dimensional +/-, and a recommended take.")

# Pillar 1: Authenticity
add_heading(doc, "Pillar 1 — Authenticity", 2)
add_styled_para(doc, "Bank's question: \"Is this photo real, or is it fabricated, edited, or AI-generated?\"", italic=True)
add_styled_para(doc, "Six items: Photo recycling detection (1.1), Cross-photo consistency (1.2), Edit/Photoshop forgery via ELA (1.3), AI-generated photo detection (1.4), Trusted-source overlay validation (1.5), and Low-tech consumer-app fraud detection (1.6 — NEW).")

add_eyebrow(doc, "Item 1.6 (NEW) — highest-leverage addition")
add_heading(doc, "Low-Tech Indonesian Consumer-App Fraud Detection", 3)
add_styled_para(doc, "Threat: The dominant Indonesian fraud profile per Adrian's brainstorm. Specific named attacks: PDF-Word-export-reedit-reexport · Instagram Story / IG Layout overlay · Screenshot-of-photo / photo-of-screen · Photo-of-printed-photo · WhatsApp re-share / forward washing · Basic phone-editor manipulation (PicsArt, Canva).")
add_styled_para(doc, "Threat motivation: Highest. Low technical barrier × high economic incentive = the canonical Indonesian fraud profile.")
add_styled_para(doc, "Solutions:")
add_bullet(doc, "A. JPEG quantization-table fingerprinting — catches Word + phone-app re-encodes. Pitch line: \"We can detect when a photo was last saved by Microsoft Word, Photoshop, Instagram, or other consumer apps.\"")
add_bullet(doc, "B. IG-Story artifact detection — catches the Instagram Story overlay trick. Pitch: \"We catch the Instagram Story overlay trick.\"")
add_bullet(doc, "C. Screenshot detection — image dimensions match phone screen exactly = flagged. Cost zero.")
add_bullet(doc, "D. Photo-of-screen detection — moiré pattern + reflection analysis.")
add_bullet(doc, "E. Photo-of-print detection — paper texture + color cast.")
add_bullet(doc, "F. WhatsApp re-share detection — compression signature + EXIF stripping.")
add_styled_para(doc, "Take: A + B + C + D should be Phase 0 priority.", bold=True, color=SLATE_900)

# Pillar 2: Identity
add_heading(doc, "Pillar 2 — Identity", 2)
add_styled_para(doc, "Bank's question: \"Is the person submitting this who they claim to be?\"", italic=True)
add_styled_para(doc, "Five items: Phone-number binding (2.1, shipped), KTP-bound capture (2.2, Pak Jun #10), Voice-print of Pelaksana (2.3, NEW), Liveness check (2.4, NEW), Device attestation (2.5, NEW).")

add_eyebrow(doc, "Item 2.2 — Pak Jun #10")
add_heading(doc, "KTP-Bound Capture", 3)
add_styled_para(doc, "Pak Jun's exact framing (April 25): \"When there is a legal issue, he is the one responsible.\"", italic=True)
add_styled_para(doc, "Solutions:")
add_bullet(doc, "A. One-time KTP enrollment at Pelaksana onboarding (developer-side setup task).")
add_bullet(doc, "B. Per-session KTP holding photo (Pak Jun's exact ask).")
add_bullet(doc, "C. Government Dukcapil API integration (e-KTP verification).")
add_bullet(doc, "D. KTP + selfie + liveness combo (cross-listed with Item 2.4).")
add_bullet(doc, "E. KYC partnership path (Adrian's contribution) — partner with Privy/Sumsub Indonesia/Tilaka/Veriff. Don't build IDV in-house. \"A product within a product.\"")

# Pillar 3: Place
add_heading(doc, "Pillar 3 — Place", 2)
add_styled_para(doc, "Bank's question: \"Is this the actual unit it claims to be?\"", italic=True)
add_styled_para(doc, "The Indonesian-reality reframe (Adrian): AutoKon doesn't need surveyor-grade GPS. AutoKon needs (1) each unit's photos to cluster at a distinct GPS centroid, (2) the pattern of centroids to match the siteplan's relative layout, (3) new photos to fall on or near predicted positions. Pattern beats accuracy.", color=SLATE_900)
add_styled_para(doc, "Seven items: Photo GPS consistency (3.1, reframed), Video site presence (3.2, Dickson L0a), Video-photo GPS cross-reference (3.3, Dickson L4), Pre-construction baseline + drone reference (3.4, REVIVED), Project-level geo-tag (3.5, Pak Jun #11), Satellite imagery cross-check (3.6, NEW), Siteplan-relative cross-reference (3.7, NEW).")

# Pillar 4: Time
add_heading(doc, "Pillar 4 — Time", 2)
add_styled_para(doc, "Bank's question: \"Was this captured now, not last week or recycled?\"", italic=True)
add_styled_para(doc, "Threat-motivation re-tier (per Adrian): Within-day time fraud is low-motivation for developers. Real time-related developer fraud is week-scale replay and recycling. Sun-angle and call-to-prayer catch sub-hour shifts that developers don't actually commit — keep for pitch theater (memorable in demos), deprioritize for engineering.")
add_table(doc,
    headers=["Tier", "Items", "Why"],
    rows=[
        ("A (ship hard)", "4.1 Audio fingerprint · 4.2 Freshness · 4.5 Public sealing", "Catches time fraud developers actually commit"),
        ("B (ship lighter)", "4.4 Behavioral baseline", "Catches Pelaksana side-gigs"),
        ("C (pitch theater)", "4.3 Sun-angle · 4.6 Call-to-prayer · 4.6 BMKG", "Memorable demos. Developers don't have motive."),
    ],
    col_widths=[1.5, 3.5, 1.5])

# Pillar 5: Chain of Custody
add_heading(doc, "Pillar 5 — Chain of Custody", 2)
add_styled_para(doc, "Bank's question: \"Once evidence enters AutoKon, can it be tampered with?\"", italic=True)
add_styled_para(doc, "Adrian's calibration: Pillar 5 is the audit-credibility narrative, not volume defense. Banks audit chain-of-custody during due diligence and incident reviews. Most Indonesian fraud is low-tech (Pillar 1 Item 1.6); Pillar 5 protects against the rare-but-catastrophic.")
add_styled_para(doc, "Six items: Multi-step approval (5.1, shipped), Tamper-evident hash chain (5.2, NEW), Periodic public sealing (5.3), SM digital sign-off (5.4, Dickson Q5), Cryptographic per-signer signing (5.5, NEW), Immutable audit log (5.6, NEW).")

# Pillar 6: Independent Verification
add_heading(doc, "Pillar 6 — Independent Verification", 2)
add_styled_para(doc, "Bank's question: \"Do you ever rely on something other than your own software?\"", italic=True)
add_styled_para(doc, "Five items: KJPP physical inspection (6.1, Indonesian standard), RT/RW or Lurah witness signature (6.2, REVIVED — uniquely Indonesian), Satellite imagery cross-check (6.3, cross-listed), Drone overflight (6.4, cross-listed), Independent third-party signing (6.5, NEW).")

doc.add_page_break()

# === Section 4: Master Item Table ===
add_heading(doc, "§4 Master Item Table", 1)
add_styled_para(doc, "Every item across 6 pillars in a single table.")
add_table(doc,
    headers=["#", "Item", "Pillar", "Phase", "Status"],
    rows=[
        ("1.1", "Photo recycling detection", "Auth", "P0", "Shipped"),
        ("1.2", "Cross-photo consistency", "Auth", "P0", ""),
        ("1.3", "Edit/Photoshop forgery (ELA)", "Auth", "P0", ""),
        ("1.4", "AI-generated detection", "Auth", "P0", ""),
        ("1.5", "Trusted-source overlay", "Auth", "P0", "Shipped"),
        ("1.6", "Low-tech consumer-app fraud", "Auth", "P0", "NEW"),
        ("2.1", "Phone-number binding", "Identity", "P0", "Shipped"),
        ("2.2", "KTP-bound capture", "Identity", "P1", ""),
        ("2.3", "Voice-print Pelaksana", "Identity", "P0", "NEW"),
        ("2.4", "Liveness check", "Identity", "P1", "NEW"),
        ("2.5", "Device attestation", "Identity", "P2", "NEW"),
        ("3.1", "Photo GPS consistency", "Place", "P0", ""),
        ("3.2", "Video site presence", "Place", "P0", "Shipped"),
        ("3.3", "Video-photo GPS cross-ref", "Place", "P0", "Shipped"),
        ("3.4", "Pre-construction baseline + drone", "Place", "P0", "Revived"),
        ("3.5", "Project-level geo-tag", "Place", "P0", ""),
        ("3.6", "Satellite imagery cross-check", "Place", "P0", "NEW"),
        ("3.7", "Siteplan-relative cross-ref", "Place", "P0", "NEW"),
        ("4.1", "Audio fingerprint", "Time", "P0", "Shipped"),
        ("4.2", "Photo freshness + monotonic", "Time", "P0", "Shipped"),
        ("4.3", "Sun-angle / shadow", "Time", "P0", ""),
        ("4.4", "Behavioral baseline", "Time", "P0", ""),
        ("4.5", "External time anchor", "Time", "P1", "NEW"),
        ("4.6", "Indonesian-context", "Time", "P1", "NEW"),
        ("5.1", "Multi-step approval", "Custody", "P0", "Shipped"),
        ("5.2", "Tamper-evident hash chain", "Custody", "P0", "NEW"),
        ("5.3", "Periodic public sealing", "Custody", "P1", "NEW"),
        ("5.4", "SM digital sign-off", "Custody", "P0", ""),
        ("5.5", "Crypto per-signer signing", "Custody", "P1", "NEW"),
        ("5.6", "Immutable audit log", "Custody", "P0", "NEW"),
        ("6.1", "KJPP physical inspection", "Indep", "P0", ""),
        ("6.2", "RT/RW witness signature", "Indep", "P1", "Revived"),
        ("6.5", "Independent third-party signing", "Indep", "P1", "NEW"),
    ])

doc.add_page_break()

# === Section 5: Recommended Ship Priority ===
add_heading(doc, "§5 Recommended Ship Priority", 1)

add_heading(doc, "Phase 0 — Ship Now", 2)
add_styled_para(doc, "Highest-leverage immediate wins:")
add_bullet(doc, "1.6 A-D — Low-tech consumer-app fraud detectors (JPEG quantization, IG-Story, screenshot, photo-of-screen). Highest-leverage new addition.")
add_bullet(doc, "5.2 D — WORM storage config change. Ships in a day. Highest cost-effectiveness ratio in the brief.")
add_bullet(doc, "5.4 A — Botpress witness yes/no prompt. 10-minute flow change.")
add_bullet(doc, "3.7 A + B — Centroid-pattern verification + siteplan-based prediction. Adrian's pattern insight applied.")
add_bullet(doc, "4.1 B — Spectral fingerprint (Shazam-style) hardening upgrade to existing audio fingerprint.")
add_bullet(doc, "6.1 B — AutoKon-routed KJPP triggers using existing tier classification.")

add_heading(doc, "Phase 1 — PILOT-30 Commitments", 2)
add_styled_para(doc, "Ship within 30 days of first signed bank pilot.")
add_bullet(doc, "Identity: 2.2 B+E — Per-session KTP + KYC partnership · 2.4 A+E — Passive liveness + WhatsApp-native KYC")
add_bullet(doc, "Place: 3.4 B — Drone fly-over · 3.6 B — Planet Labs")
add_bullet(doc, "Time: 4.5 B — RFC 3161 TSA · 4.6 A+B — BMKG weather + call-to-prayer (Indonesian moats) · 4.1 D+E — Multi-segment + acoustic environment")
add_bullet(doc, "Custody: 5.2 A+B — Per-project Merkle + per-unit hash chain · 5.5 A+C — Server signing + PKI public registry")
add_bullet(doc, "Independent: 6.2 A — RT/RW witness · 6.5 — Privy/eMeterai third-party signing")

add_heading(doc, "Phase 2 — Roadmap", 2)
add_bullet(doc, "1.4 B — C2PA content credentials (waits on phone fleet upgrade ~2027+)")
add_bullet(doc, "2.5 A — Android Play Integrity (requires custom AutoKon camera app)")
add_bullet(doc, "1.1 C — PRNU sensor fingerprint (theoretical for Indonesian field)")
add_bullet(doc, "5.5 B+D — Mobile-device-bound keys + HSM keys")
add_bullet(doc, "5.6 C — Database WAL export to immutable storage")

add_heading(doc, "Roadmap Reserve — Defer Until Observed", 2)
add_bullet(doc, "Advanced AI-generation defeats — until real attack observed")
add_bullet(doc, "VFX video forgery — until observed")
add_bullet(doc, "Cross-modality consistent forgery — until observed")
add_bullet(doc, "Rooted device sophisticated bypass — until observed at scale")

doc.add_page_break()

# === Section 6: Pitch Language ===
add_heading(doc, "§6 Pitch Language Reference", 1)
add_styled_para(doc, "Actual sentences for the bank pitch deck, organized by pillar.")

pitch = [
    ("Authenticity", [
        "We screen every photo against state-of-the-art AI-generation detectors — the same techniques used to catch the JAKI Pemprov DKI Gemini-watermark scandal in 2025.",
        "Our audio fingerprint catches recycled videos using the same forensic technique Shazam uses to identify songs.",
        "We catch the Microsoft-Word-export trick. We catch the Instagram Story overlay trick. We catch screenshots. We catch photos of printed photos. We've engineered for what Indonesian fraudsters actually do.",
    ]),
    ("Identity", [
        "Every Pelaksana is KTP-bound. We use the same KYC vendor your bank's account-opening team uses.",
        "Voice-print verification at every session.",
        "Bank-PIC signatures go through Privy / eMeterai — the same e-signature infrastructure your KYC team uses.",
    ]),
    ("Place", [
        "We don't claim perfect GPS coordinates — Indonesia's archipelago infrastructure makes that impossible. We claim 1000 units have 1000 distinct GPS signatures matching the developer-provided siteplan layout.",
        "We integrate the developer's existing drone monitoring as independent aerial cross-reference.",
        "Cross-checked against EU Space Agency Sentinel-2 satellite imagery.",
        "Project location is locked at bank onboarding and continuously verified against field-captured unit GPS.",
    ]),
    ("Time", [
        "We use astronomy to verify your photo was taken at the time you claim — sun angle, shadow direction, calculated from GPS and timestamp.",
        "If the call to prayer is audible in the walk-around video, our system verifies it matches the muezzin schedule for that location.",
        "Cross-checked against BMKG's official weather records.",
        "Every day, our evidence hash is sealed to an RFC 3161 Time-Stamping Authority — the same standard used for legal e-signatures.",
    ]),
    ("Chain of Custody", [
        "Every record in AutoKon has a complete edit history, never a silent overwrite. Every project has a Merkle tree of all evidence — any tampering becomes detectable.",
        "All evidence files are stored on Write-Once-Read-Many infrastructure. Even AutoKon admins cannot delete or overwrite evidence.",
        "Site managers must explicitly attest, in Bahasa, that they personally witnessed the work.",
    ]),
    ("Independent Verification", [
        "AutoKon makes KJPP 10x more focused. Software pre-screens; KJPP inspects what software flags as risky.",
        "Local RT/RW officials witness major milestones — independent verification rooted in Indonesian community structure.",
        "We don't ask the bank to trust only us — Privy/eMeterai holds an independent record of every signature.",
    ]),
]
for pillar, lines in pitch:
    add_heading(doc, pillar, 3)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('"' + line + '"')
        run.italic = True
        run.font.name = "Inter"
        run.font.size = Pt(11)
        run.font.color.rgb = SLATE_700

doc.add_page_break()

# === Section 7: Open Questions ===
add_heading(doc, "§7 Open Questions for Decision", 1)
add_styled_para(doc, "Eight decisions before engineering can proceed cleanly.")
opens = [
    "Is custom AutoKon camera app on the table for PILOT-30? Cross-cuts Items 1.5B + 2.5A. Resolves device-attestation and overlay-spoofing simultaneously. Premium-tier feature.",
    "KYC partnership selection. Privy vs Tilaka vs Sumsub Indonesia vs Veriff. Commercial relationship + integration choice.",
    "Sentinel-2 free vs Planet Labs paid for Phase 0 satellite imagery.",
    "OpenTimestamps (Bitcoin/Ethereum) vs RFC 3161 TSA for public sealing language preference.",
    "WORM storage migration — Supabase storage config change to immutability mode.",
    "Drone overflight ingest format (MP4? GeoTIFF?) standardization with developers.",
    "Tier-3 review queue SLA (Dickson's open question 6).",
    "RT/RW witness program scope. Pilot first or default for all bank-pitch tier? Honoraria budget per project.",
]
for o in opens:
    add_numbered(doc, o)

# === Section 8: Appendix ===
add_heading(doc, "§8 Appendix — Relationship to Existing AutoKon Work", 1)
add_table(doc,
    headers=["Existing work", "This brief's relationship"],
    rows=[
        ("Dickson's evidence integrity summary (May 1)", "Reorganized. 10-layer framework preserved underneath; selectively-revived rejected items added per Adrian's R&D push-back."),
        ("Pak Jun product recommendations (Apr 24-25)", "Items #2 KodeProper · #5 trigger points · #9 digital signatures · #10 KTP capture · #11 project geo-tag · #12 RBAC — all integrated."),
        ("autokon-app production code", "No PRs proposed. R&D-mode means we list options."),
        ("autokon-bank-report-pitch repo", "Frozen. Brief is for the bank-pitch evidence-integrity narrative but doesn't modify that repo."),
        ("autokon-three-audience-samples repo", "Cross-referenced. Audience-first framing consistent across both."),
        ("V2.1 Continuous Progress Snapshot", "Foundational primitive. Some items extend CPS's tamper-resistance."),
        ("F3a Design System Spec", "Independent. Different scope; consistent voice."),
    ])

# === Footer ===
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("End of brief.")
run.italic = True
run.font.name = "Inter"
run.font.size = Pt(11)
run.font.color.rgb = SLATE_600

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AutoKon · Evidence Integrity R&D Brief · v1.0 · 2026-05-01")
run.font.name = "JetBrains Mono"
run.font.size = Pt(9)
run.font.color.rgb = SLATE_500

# Save
doc.save("AutoKon-Evidence-Integrity-Research.docx")
print("Saved: AutoKon-Evidence-Integrity-Research.docx")
